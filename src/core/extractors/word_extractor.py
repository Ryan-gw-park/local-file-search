"""
Local Finder X v2.0 - Word Document Extractor

Extracts content from .docx files using python-docx.
"""

from typing import List, Dict, Any, Optional

from .base import BaseExtractor, ExtractorResult, register_extractor

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    PackageNotFoundError = Exception


@register_extractor
class WordExtractor(BaseExtractor):
    """Extractor for Microsoft Word documents (.docx)."""
    
    SUPPORTED_EXTENSIONS = [".docx"]
    
    def extract(self, file_path: str) -> ExtractorResult:
        """Extract content from a Word document."""
        if not DOCX_AVAILABLE:
            return self._create_error_result(
                "python-docx is not installed. Install with: pip install python-docx"
            )
        
        try:
            doc = Document(file_path)
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Extract content with structure
            sections = []
            full_text_parts = []
            current_section = {"type": "paragraph", "content": []}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if this is a heading
                if para.style and para.style.name.startswith("Heading"):
                    # Save previous section
                    if current_section["content"]:
                        sections.append(current_section)
                    
                    # Start new section with heading
                    heading_level = self._get_heading_level(para.style.name)
                    current_section = {
                        "type": "heading",
                        "level": heading_level,
                        "title": text,
                        "content": [],
                    }
                else:
                    current_section["content"].append(text)
                
                full_text_parts.append(text)
            
            # Add last section
            if current_section["content"] or current_section.get("title"):
                sections.append(current_section)
            
            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    sections.append({
                        "type": "table",
                        "content": [table_text],
                    })
                    full_text_parts.append(table_text)
            
            full_text = "\n".join(full_text_parts)
            
            return self._create_success_result(
                text=full_text,
                sections=sections,
                metadata=metadata,
            )
            
        except PackageNotFoundError:
            return self._create_error_result(f"File not found or invalid: {file_path}")
        except Exception as e:
            return self._create_error_result(f"Error extracting Word document: {str(e)}")
    
    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract document metadata."""
        metadata = {}
        
        try:
            props = doc.core_properties
            if props.author:
                metadata["author"] = props.author
            if props.title:
                metadata["title"] = props.title
            if props.subject:
                metadata["subject"] = props.subject
            if props.created:
                metadata["created"] = props.created.isoformat()
            if props.modified:
                metadata["modified"] = props.modified.isoformat()
        except Exception:
            pass
        
        return metadata
    
    def _get_heading_level(self, style_name: str) -> int:
        """Get heading level from style name."""
        try:
            # "Heading 1" -> 1
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 1
    
    def _extract_table(self, table) -> str:
        """Extract text from a table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


__all__ = ["WordExtractor", "DOCX_AVAILABLE"]
